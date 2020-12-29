import docx                                                                                             # pip install python-docx
from nltk.stem import SnowballStemmer                                                                   # pip install nltk
from nltk.corpus import stopwords
from nltk import tokenize, FreqDist
import os
import tika
TIKA_SERVER_JAR = 'https://repo1.maven.org/maven2/org/apache/tika/tika-server/1.19/tika-server-1.19.jar'
tika.TikaClientonly = True
tika.initVM()

from tika import parser, language                                                                                 # pip install tika
from langid import classify                                                                             # pip install langid
from re import sub
import pandas as pd                                                                                     # pip install pandas
from tkinter import filedialog
import tkinter as tk


def vectorize(tokens):                                                                                  # Define class vectorize
    vector = []
    for w in keywords:
        vector.append(tokens.count(w))
    return vector


language = {
    "nl": "dutch",
    "en": "english",
    "es": "spanish",
    "fr": "french",
    "de": "german",
    "it": "italian"
    }

with open("keywords.txt", "r") as keywords:                                                             # Import file with keywords
    keywords = keywords.read().split()                                                                  # Convert keywords into list
index = ['Source', 'Lang']
headings_bow = index + keywords                                                                         # Create headings for bow-table (Bag of Words)


root = tk.Tk()

canvas1 = tk.Canvas(root, width=400, height=300)                                                        # Create empty program window
canvas1.pack()


directory_path = filedialog.askdirectory()


def analysis():
    directory = os.fsencode(directory_path)
    bow_df = pd.DataFrame(columns=headings_bow)                                                         # Create empty table for bow
    report = docx.Document()                                                                            # Create report document
    report.add_heading(f'Analysis {os.path.basename(directory_path)}', 0)                               # Add title to report
    for file in os.listdir(directory):
        document_path = os.path.join(directory, file).decode()
        document = parser.from_file(document_path)                                                      # Retrieve text from file
        document = document['content']
        content = sub(r'http\S+', " ", document)                                                        # Delete all links
        content = sub("[^a-zA-Z0-9|^-]", " ", content).lower()                                          # Delete all punctuation/upper case letters
        content_words = tokenize.word_tokenize(content)                                                 # Split words into list
        language_name = language[classify(content)[0]]                                                  # Detect text language
        content_words_core = " ".join(filter(lambda x: x in keywords, content_words)).split()           # Delete all words except for words in keywords
        vector = vectorize(content_words_core)                                                          # Count occurrence of keywords
        filename_first = os.fsdecode(file)[0:3]                                                         # Select first 3 characters of filename
        vector.insert(0, language_name.capitalize())                                                    # Add language to vector-list
        vector.insert(0, filename_first)                                                                # Add first 3 characters of filename to vector-list
        bow = pd.DataFrame(vector).transpose()                                                          # Put vector-list into table and transpose
        bow.columns = headings_bow                                                                      # Add headings to table
        bow_df = pd.concat([bow_df, bow])                                                               # Add table to table of all files
        bow_df[keywords] = bow_df[keywords].astype('int64')                                             # Change datatype in table to integer
    bow_df.loc[:, 'Total'] = bow_df.sum(numeric_only=True, axis=1)                                      # Add totals column
    bow_df.sort_values(by=['Total'], inplace=True, ascending=False)                                     # Sort table on descending total column
    table_bow = report.add_table(bow_df.shape[0] + 1, bow_df.shape[1])                                  # Add template table
    for j in range(bow_df.shape[-1]):
        table_bow.cell(0, j).text = bow_df.columns[j]                                                   # Add headers to table
    for i in range(bow_df.shape[0]):
        for j in range(bow_df.shape[-1]):
            table_bow.cell(i + 1, j).text = str(bow_df.values[i, j])                                    # Add data to table
    table_bow.style = 'Light Shading'                                                                   # Change style of table

    for file in os.listdir(directory):
        document_path = os.path.join(directory, file).decode()
        document = parser.from_file(document_path)                                                      # Retrieve text from file
        document = document['content']
        content = sub(r'http\S+', " ", document)                                                        # Delete all links
        content = sub("[^a-zA-Z|^-]", " ", content).lower()                                             # Delete all punctuation/upper case letters/numbers
        content_words = [w for w in content.split() if len(w) > 1]                                      # Delete all words with one letter and split words into list
        language_name = language[classify(content)[0]]                                                  # Detect text language
        content_words_core = [w for w in content_words if w not in stopwords.words(language_name)]      # Delete adverbs
        stemmed_words = [SnowballStemmer(language_name).stem(word) for word in content_words_core]      # Group different forms of a word to a single item
        fdist1 = FreqDist(stemmed_words)                                                                # Count occurrence of words
        top_10_words = pd.DataFrame(fdist1.most_common(10), columns=['Word', 'Count'])                  # Put top 10 words in table
        filename = os.fsdecode(file)                                                                    # Retrieve filename
        report.add_heading(filename, level=1)                                                           # Add subtitle per document
        report.add_paragraph(f'Language: {language_name.capitalize()}')                                 # Add language
        table = report.add_table(top_10_words.shape[0] + 1, top_10_words.shape[1])                      # Add template table
        for j in range(top_10_words.shape[-1]):
            table.cell(0, j).text = top_10_words.columns[j]                                             # Add headers to table
        for i in range(top_10_words.shape[0]):
            for j in range(top_10_words.shape[-1]):
                table.cell(i + 1, j).text = str(top_10_words.values[i, j])                              # Add data to table
        table.style = 'Light Shading'                                                                   # Change style of table

    report.save(f'{os.environ["USERPROFILE"]}/Desktop/report.docx')                                     # Save document to desktop


button1 = tk.Button(text='Start analysis!', command=analysis)                                           # Create button executing Analysis
canvas1.create_window(200, 180, window=button1)

root.mainloop()
